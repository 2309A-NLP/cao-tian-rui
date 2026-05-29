# -*- coding: utf-8 -*-
"""
数据清理和分块模块 - 支持动态分块策略，集成高级文档解析器

⚠️ 常改动的地方：
1. 分块策略参数（chunk_size, chunk_overlap, min_chunk_size）在 ChunkStrategy.STRATEGIES 中调整
2. 关键词映射（KEYWORD_MAPPING）用于识别文件类型并匹配分块策略
3. 清理器中的正则表达式规则（如移除特殊字符、重复标点等）
4. 过滤文本的最小长度（min_length 默认 20，可在调用处或策略中修改）
5. 高级解析器的启用/禁用（enable_ocr, use_gpu 参数）

⚠️ 注意事项：
1. 高级解析器（AdvancedDocumentParser）优先用于 PDF/DOCX，失败时降级到基础解析器
2. 文档处理流程：原始文本 → 清洗 → 过滤 → 分块 → 生成 Document 对象
3. 分块策略根据文件名中的关键词自动选择（法律/医疗/理财/通用）
4. 所有处理步骤都会通过 document_tracker 保存中间产物（raw/cleaned/chunks）
5. DataProcessor 依赖 utils.document_tracker 记录处理报告
"""

import os
import re
import json
import csv
import logging
from datetime import datetime
from typing import List, Dict, Any, Optional, Tuple
from pathlib import Path
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter

from utils.document_tracker import document_tracker
from utils.logger import get_logger

# 高级解析器（优先使用）
try:
    from .advanced_parser import AdvancedDocumentParser
    HAS_ADVANCED_PARSER = True
    print("[OK] 高级文档解析器已加载")
except ImportError:
    HAS_ADVANCED_PARSER = False
    print("[信息] 高级文档解析器未安装，使用基础解析")

# 可选依赖：PDF 解析
try:
    from PyPDF2 import PdfReader
    HAS_PDF = True
except ImportError:
    HAS_PDF = False
    print("警告: PyPDF2 未安装，无法解析 PDF 文件")

# 可选依赖：DOCX 解析
try:
    from docx import Document as DocxDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    print("警告: python-docx 未安装，无法解析 DOCX 文件")

# 配置日志（同时输出到文件和控制台）
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('data_processing.log', encoding='utf-8'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)


class DataCleaner:
    """数据清理器：清洗文本中的噪音字符、统一标点、去除多余空白"""

    @staticmethod
    def clean_text(text: str) -> str:
        """
        清洗文本
        ⚠️ 常改动：正则表达式规则可根据实际文档格式调整
        """
        if not text:
            return ""

        original_length = len(text)

        # 1. 移除特殊字符（保留中英文、数字、常用标点和数学运算符）
        # 允许保留：中文字符、大小写字母、数字、空白、常见标点及运算符号
        text = re.sub(r'[^\u4e00-\u9fa5a-zA-Z0-9\s\.,!?:;。，：；？！""''《》【】（）+\-*/=<>]', '', text)

        # 2. 移除多余空白（换行、制表、连续空格替换为单个空格）
        text = re.sub(r'\s+', ' ', text)

        # 3. 移除首尾空格
        text = text.strip()

        # 4. 移除重复的标点（如多个句号/感叹号/问号/分号合并为一个）
        text = re.sub(r'([。！？；])\1+', r'\1', text)

        # 5. 统一标点符号为英文格式（可选，可根据需求关闭）
        text = text.replace('，', ',').replace('。', '.').replace('！', '!').replace('？', '?')

        cleaned_length = len(text)

        if original_length > cleaned_length:
            logger.debug(f"文本清理: {original_length} -> {cleaned_length} 字符")

        return text

    @staticmethod
    def filter_valid_texts(texts: List[str], min_length: int = 20) -> List[str]:
        """
        过滤掉过短或无效的文本段
        ⚠️ 常改动：min_length 默认 20，可根据文档类型调整（法律可更大）
        """
        filtered = []
        for text in texts:
            cleaned = text.strip()
            if len(cleaned) >= min_length:
                filtered.append(cleaned)

        removed_count = len(texts) - len(filtered)
        if removed_count > 0:
            logger.info(f"过滤无效文本: 移除 {removed_count} 段 (小于{min_length}字符)")

        return filtered


class DocumentLoader:
    """
    文档加载器 - 支持多种格式（TXT, MD, PDF, DOCX, JSON, CSV）
    优先使用高级解析器（AdvancedDocumentParser）处理 PDF/DOCX，降级采用基础库
    """

    def __init__(self):
        self.supported_formats = {
            '.txt': self.load_txt,
            '.md': self.load_txt,
            '.pdf': self.load_pdf,
            '.docx': self.load_docx,
            '.json': self.load_json,
            '.csv': self.load_csv,
        }
        # 初始化高级解析器（如果可用）
        self.advanced_parser = None
        if HAS_ADVANCED_PARSER:
            try:
                # ⚠️ 常改动：enable_ocr 和 use_gpu 可配置
                self.advanced_parser = AdvancedDocumentParser(enable_ocr=False, use_gpu=False)
                print("[OK] 高级解析器初始化成功")
            except Exception as e:
                print(f"[警告] 高级解析器初始化失败: {e}")

    def load_document(self, file_path: str) -> List[str]:
        """加载单个文档，返回文本段列表"""
        file_ext = Path(file_path).suffix.lower()

        if file_ext not in self.supported_formats:
            logger.warning(f"不支持的文件格式: {file_ext} - {file_path}")
            return []

        loader = self.supported_formats[file_ext]
        if loader is None:
            logger.warning(f"缺少解析库: {file_ext} - {file_path}")
            return []

        try:
            logger.info(f"加载文件: {Path(file_path).name}")
            texts = loader(file_path)
            logger.info(f"  提取 {len(texts)} 段文本")
            return texts
        except Exception as e:
            logger.error(f"加载失败 {file_path}: {e}")
            return []

    def load_directory(self, directory_path: str, recursive: bool = True) -> Dict[str, List[str]]:
        """加载目录下所有支持格式的文档，返回 {文件路径: 文本段列表}"""
        all_texts = {}
        directory = Path(directory_path)

        if not directory.exists():
            logger.error(f"目录不存在: {directory_path}")
            return all_texts

        pattern = "**/*" if recursive else "*"

        for file_path in directory.glob(pattern):
            if file_path.is_file() and file_path.suffix.lower() in self.supported_formats:
                texts = self.load_document(str(file_path))
                if texts:
                    all_texts[str(file_path)] = texts

        logger.info(f"目录加载完成: 共加载 {len(all_texts)} 个文件")
        return all_texts

    def load_txt(self, file_path: str) -> List[str]:
        """加载 TXT/MD 文件：按双换行分割段落，否则整体返回"""
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        paragraphs = [p.strip() for p in content.split('\n\n') if p.strip()]
        return paragraphs if paragraphs else [content]

    def load_pdf(self, file_path: str) -> List[str]:
        """
        加载 PDF 文件：优先使用高级解析器（可提取表格、保留页码），降级使用 PyPDF2
        """
        # 尝试使用高级解析器
        if self.advanced_parser is not None:
            try:
                print(f"  [高级解析] 正在解析: {Path(file_path).name}")
                result = self.advanced_parser.parse_pdf_deep(file_path)

                texts = []
                for page in result["pages"]:
                    if page["text"].strip():
                        texts.append(f"[第{page['page_num']}页]\n{page['text']}")

                for table in result["tables"]:
                    if table["text"].strip():
                        texts.append(f"[表格 第{table['page']}页]\n{table['text']}")

                print(f"  [高级解析] 提取 {len(texts)} 段文本")
                if texts:
                    return texts

            except Exception as e:
                print(f"  [高级解析] 失败: {e}，降级使用 PyPDF2")

        if not HAS_PDF:
            print("  [警告] PyPDF2 未安装，无法解析 PDF")
            return []

        try:
            reader = PdfReader(file_path)
            texts = []
            for page_num, page in enumerate(reader.pages, 1):
                text = page.extract_text()
                if text and text.strip():
                    texts.append(f"[第{page_num}页]\n{text.strip()}")
            return texts
        except Exception as e:
            logger.error(f"PDF加载失败 {file_path}: {e}")
            return []

    def load_docx(self, file_path: str) -> List[str]:
        """
        加载 Word 文档：优先使用高级解析器，降级使用 python-docx
        """
        if self.advanced_parser is not None:
            try:
                print(f"  [高级解析] 正在解析DOCX: {Path(file_path).name}")
                result = self.advanced_parser.parse_pdf_deep(file_path)

                texts = []
                for page in result["pages"]:
                    if page["text"].strip():
                        texts.append(f"[第{page['page_num']}页]\n{page['text']}")

                for table in result["tables"]:
                    if table["text"].strip():
                        texts.append(f"[表格 第{table['page']}页]\n{table['text']}")

                if texts:
                    print(f"  [高级解析] 提取 {len(texts)} 段文本")
                    return texts

            except Exception as e:
                print(f"  [高级解析] 失败: {e}，降级使用 python-docx")

        if not HAS_DOCX:
            print("  [警告] python-docx 未安装，无法解析 DOCX")
            return []

        try:
            doc = DocxDocument(file_path)
            texts = []

            for para in doc.paragraphs:
                if para.text.strip():
                    texts.append(para.text.strip())

            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                    if row_text:
                        texts.append(row_text)

            return texts
        except Exception as e:
            logger.error(f"DOCX加载失败 {file_path}: {e}")
            return []

    def load_json(self, file_path: str) -> List[str]:
        """加载 JSON 文件：递归提取所有字符串值（长度>10）作为独立文本段"""
        with open(file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)

        texts = []

        def extract_text(obj, prefix=""):
            if isinstance(obj, dict):
                for key, value in obj.items():
                    if isinstance(value, str) and len(value) > 10:
                        texts.append(f"{prefix}{key}: {value}")
                    elif isinstance(value, (dict, list)):
                        extract_text(value, f"{prefix}{key}.")
            elif isinstance(obj, list):
                for i, item in enumerate(obj):
                    if isinstance(item, str) and len(item) > 10:
                        texts.append(f"{prefix}[{i}]: {item}")
                    elif isinstance(item, (dict, list)):
                        extract_text(item, f"{prefix}[{i}].")

        extract_text(data)
        return texts

    def load_csv(self, file_path: str) -> List[str]:
        """加载 CSV 文件：将每行的键值对组合成一个文本段"""
        texts = []
        with open(file_path, 'r', encoding='utf-8') as f:
            reader = csv.DictReader(f)
            for row in reader:
                row_text = '，'.join([f"{k}:{v}" for k, v in row.items() if v and v.strip()])
                if row_text:
                    texts.append(row_text)
        return texts


class ChunkStrategy:
    """分块策略配置，根据文件内容类型（法律、医疗、理财）选择不同的分块参数"""

    # 预定义策略
    STRATEGIES = {
        "law": {
            "chunk_size": 900,
            "chunk_overlap": 150,
            "min_chunk_size": 20,
            "separators": ['\n\n', '。', '；', '\n', '，', ' ', ''],
            "description": "法律条文策略"
        },
        "medical": {
            "chunk_size": 700,
            "chunk_overlap": 100,
            "min_chunk_size": 100,
            "separators": ['\n\n', '\n', '。', '！', '？', '；', '，', ' ', ''],
            "description": "医疗指南策略"
        },
        "finance": {
            "chunk_size": 600,
            "chunk_overlap": 80,
            "min_chunk_size": 80,
            "separators": ['\n\n', '\n', '。', '！', '？', '；', '，', ' ', ''],
            "description": "理财文档策略"
        },
        "default": {
            "chunk_size": 800,
            "chunk_overlap": 100,
            "min_chunk_size": 50,
            "separators": ['\n\n', '\n', '。', '！', '？', '；', '，', ' ', ''],
            "description": "通用策略"
        }
    }

    # 关键词映射：根据文件名中包含的关键词匹配策略
    # ⚠️ 常改动：可增加更多关键词或调整匹配优先级
    KEYWORD_MAPPING = {
        "law": ["法律", "民法典", "法规", "条文", "司法", "判决", "law", "legal"],
        "medical": ["医疗", "医学", "临床", "指南", "诊疗", "疾病", "高血压", "心衰", "medical", "health"],
        "finance": ["理财", "金融", "投资", "股票", "基金", "finance", "investment"],
    }

    @classmethod
    def get_strategy(cls, file_path: str) -> Dict:
        """根据文件路径（文件名）选择最匹配的分块策略"""
        path_lower = file_path.lower()

        for strategy_type, keywords in cls.KEYWORD_MAPPING.items():
            for keyword in keywords:
                if keyword in path_lower:
                    strategy = cls.STRATEGIES.get(strategy_type, cls.STRATEGIES["default"])
                    logger.info(f"为文件 {Path(file_path).name} 选择策略: {strategy['description']}")
                    return strategy

        return cls.STRATEGIES["default"]


class AdaptiveTextSplitter:
    """自适应文本分块器：根据文件类型创建合适的 RecursiveCharacterTextSplitter"""

    def __init__(self, default_chunk_size: int = 800, default_chunk_overlap: int = 100):
        self.default_chunk_size = default_chunk_size
        self.default_chunk_overlap = default_chunk_overlap

    def get_splitter_for_file(self, file_path: str) -> Tuple[RecursiveCharacterTextSplitter, Dict]:
        """为特定文件创建分块器，返回 (splitter, strategy)"""
        strategy = ChunkStrategy.get_strategy(file_path)

        splitter = RecursiveCharacterTextSplitter(
            chunk_size=strategy["chunk_size"],
            chunk_overlap=strategy["chunk_overlap"],
            separators=strategy["separators"],
            length_function=len,
        )

        logger.info(f"  分块参数: chunk_size={strategy['chunk_size']}, overlap={strategy['chunk_overlap']}")

        return splitter, strategy


class DataProcessor:
    """
    完整的数据处理器：整合加载、清洗、过滤、分块、追踪和报告
    ⚠️ 常改动：process_directory 中的递归参数，默认 True
    """

    def __init__(self, chunk_size: int = 800, chunk_overlap: int = 100):
        self.cleaner = DataCleaner()
        self.loader = DocumentLoader()
        self.adaptive_splitter = AdaptiveTextSplitter(chunk_size, chunk_overlap)
        self.processing_stats = {}
        self.all_strategy_stats = {}
        self.rag_logger = get_logger("rag")

    def process_directory(self, directory_path: str) -> List[Document]:
        """处理整个目录下的所有文档，返回 Document 列表（用于向量库）"""
        all_documents = []

        files_data = self.loader.load_directory(directory_path)

        for file_path, texts in files_data.items():
            logger.info(f"\n处理文件: {Path(file_path).name}")
            documents = self._process_single_file(file_path, texts)
            all_documents.extend(documents)

            strategy = ChunkStrategy.get_strategy(file_path)
            strategy_name = strategy.get("description", "unknown")
            self.all_strategy_stats[strategy_name] = self.all_strategy_stats.get(strategy_name, 0) + 1

        logger.info(f"\n所有文件处理完成: 共 {len(all_documents)} 个块")

        if all_documents:
            self.processing_stats = {
                "total_chunks": len(all_documents),
                "total_chars": sum(len(doc.page_content) for doc in all_documents),
                "source_files": len(files_data)
            }
            # 保存完整处理报告到 data/processed_docs/processing_reports.json
            document_tracker.save_processing_report(directory_path, self.processing_stats)

        return all_documents

    def _process_single_file(self, file_path: str, texts: List[str]) -> List[Document]:
        """
        处理单个文件：保存原始文本 → 清洗 → 保存清洗后文本 → 过滤 → 分块 → 保存分块结果
        ⚠️ 注意事项：分块后会为每个 chunk 添加索引元数据
        """
        import time
        start_time = time.time()

        strategy = ChunkStrategy.get_strategy(file_path)
        file_name = Path(file_path).name

        # 1. 保存原始文本（用于追溯）
        document_tracker.save_raw_text(file_name, texts)

        # 2. 清洗
        cleaned_texts = []
        for text in texts:
            cleaned = self.cleaner.clean_text(text)
            if cleaned:
                cleaned_texts.append(cleaned)

        # 3. 保存清洗后文本
        document_tracker.save_cleaned_text(file_name, cleaned_texts, {
            "original_count": len(texts),
            "cleaned_count": len(cleaned_texts),
            "min_length": strategy.get("min_chunk_size", 50)
        })

        # 4. 过滤（按策略中的最小长度）
        min_length = strategy.get("min_chunk_size", 50)
        filtered_texts = self.cleaner.filter_valid_texts(cleaned_texts, min_length=min_length)

        # 5. 分割：先将每个文本段转为 Document，再用 splitter 分块
        splitter, _ = self.adaptive_splitter.get_splitter_for_file(file_path)

        documents = []
        for text in filtered_texts:
            doc = Document(
                page_content=text,
                metadata={
                    "source": file_name,
                    "source_file": file_name,
                    "file_name": file_name,
                    "source_type": "document",
                    "strategy": strategy.get("description", "default"),
                    "processed_at": datetime.now().isoformat()
                }
            )
            documents.append(doc)

        split_docs = splitter.split_documents(documents)

        # 6. 保存分块结果（JSON）
        document_tracker.save_chunks(file_name, split_docs, strategy.get("description"))

        # 7. 记录日志和统计
        elapsed = (time.time() - start_time) * 1000
        self.rag_logger.info(
            f"文件处理: {file_name} | {len(texts)}段 -> {len(split_docs)}块 | 策略: {strategy.get('description')} | 耗时: {elapsed:.0f}ms")

        logger.info(f"  文件处理完成: {len(texts)} 段 -> {len(split_docs)} 个块")

        # 为每个分块添加索引（便于调试）
        for i, doc in enumerate(split_docs):
            doc.metadata["chunk_index"] = i
            doc.metadata["chunk_total"] = len(split_docs)

        return split_docs

    def save_preview(self, documents: List[Document], output_file: str = "processed_data_preview.json"):
        """
        保存数据预览文件（用于前端展示或调试）
        ⚠️ 常改动：output_file 默认名称
        """
        preview = {
            "statistics": self.processing_stats,
            "samples": [],
            "all_chunks_summary": []
        }

        # 前10个块的详细预览
        for i, doc in enumerate(documents[:10]):
            preview["samples"].append({
                "index": i,
                "content": doc.page_content[:500] + "..." if len(doc.page_content) > 500 else doc.page_content,
                "length": len(doc.page_content),
                "metadata": doc.metadata
            })

        # 前50个块的摘要
        for i, doc in enumerate(documents[:50]):
            preview["all_chunks_summary"].append({
                "index": i,
                "length": len(doc.page_content),
                "preview": doc.page_content[:100],
                "metadata": doc.metadata
            })

        with open(output_file, 'w', encoding='utf-8') as f:
            json.dump(preview, f, ensure_ascii=False, indent=2)

        logger.info(f"数据预览已保存到: {output_file}")
        print(f"\n[数据预览] 已保存到: {output_file}")

    def print_summary(self):
        """打印处理摘要（控制台输出）"""
        print("\n" + "=" * 60)
        print("数据处理摘要")
        print("=" * 60)
        print(f"总块数: {self.processing_stats.get('total_chunks', 0):,}")
        print(f"总字符: {self.processing_stats.get('total_chars', 0):,}")
        print(f"源文件数: {self.processing_stats.get('source_files', 0)}")
        print("=" * 60)

    def process_with_advanced_parser(self, file_path: str) -> List[Document]:
        """
        使用高级解析器单独处理文件（不经过常规流程），返回 Document 列表
        ⚠️ 注意事项：此方法主要用于测试或特殊场景，常规处理已自动集成高级解析器
        """
        try:
            from .advanced_parser import AdvancedDocumentParser

            parser = AdvancedDocumentParser(enable_ocr=False)
            parsed = parser.parse_pdf_deep(file_path)
            documents = parser.convert_to_documents(parsed)

            print(f"[高级解析] 成功处理 {file_path}, 生成 {len(documents)} 个文档块")
            return documents

        except ImportError:
            print(f"[高级解析] 模块未安装，使用默认解析: {file_path}")
            return []
        except Exception as e:
            print(f"[高级解析] 失败: {e}")
            return []